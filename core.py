import math
import string

from docx import Document
from docx.oxml import CT_P, CT_Tbl


def focal_point(file_name: str):
    document = Document(file_name)
    para_count = 0
    table_count = 0
    parent_elm = document._body._body
    seen_paragraphs = set()

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            paragraph = document.paragraphs[para_count]

            if paragraph._element not in seen_paragraphs:
                seen_paragraphs.add(paragraph._element)
                para_count += 1
                # We cannot yet handle editing hyperlinks
                # as there seems to be no simple way
                # to insert them back into their original
                # position after editing.
                elem = child.xpath("./w:hyperlink")
                if elem:
                    hyperlink = elem[0]
                    paragraph._p.remove(hyperlink)

                runs = paragraph.runs
                for run, next_run in zip(runs, runs[1:]):
                    split_run_into_paragraph(run, next_run, paragraph)
                if runs:
                    split_run_into_paragraph(runs[-1], None, paragraph)

        elif isinstance(child, CT_Tbl):
            table = document.tables[table_count]
            table_count += 1
            for row in table.rows:
                copy_cells(
                    row.cells,
                    seen_paragraphs=seen_paragraphs,
                )
            for column in table.columns:
                copy_cells(
                    column.cells,
                    seen_paragraphs=seen_paragraphs,
                )

    file_name, extension = file_name.split('.')
    new_file_name = f'{file_name}_edited.{extension}'
    document.save(new_file_name)


def split_run_into_paragraph(run, next_run, paragraph):

    words = run.text.rstrip().split(' ')
    if next_run is None:
        remove_space_at_end = False
    else:
        remove_space_at_end = False
        punctuation = tuple(i for i in string.punctuation) + (' ',)
        if next_run.text.startswith(punctuation) or not run.text.endswith(' '):
            remove_space_at_end = True

    if 'graphicData' not in run._r.xml and 'w:br' not in run._element.xml:
        run.clear()

    new_run = None
    for word in words:
        bold_part, regular_part = _embolden_word(word)
        for text, bold in (
                (bold_part, True),
                (regular_part + ' ', False),
        ):
            new_run = paragraph.add_run(
                text=text, style=run.style.name if run.style else '')
            new_run.bold = bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            for attr in (
                    'all_caps',
                    'complex_script',
                    'cs_bold',
                    'cs_italic',
                    'double_strike',
                    'emboss',
                    'hidden',
                    'highlight_color',
                    'imprint',
                    'math',
                    'name',
                    'no_proof',
                    'outline',
                    'rtl',
                    'shadow',
                    'size',
                    'small_caps',
                    'spec_vanish',
                    'strike',
                    'subscript',
                    'web_hidden',
            ):
                value = getattr(run.font, attr)
                setattr(new_run.font, attr, value)

            setattr(new_run.font.color, 'rgb', run.font.color.rgb)

    if remove_space_at_end:
        new_run.text = new_run.text.rstrip()


def copy_cells(cells, seen_paragraphs):
    for cell in cells:
        for cell_paragraph in cell.paragraphs:
            if cell_paragraph._element not in seen_paragraphs:
                seen_paragraphs.add(cell_paragraph._element)
                runs = cell_paragraph.runs
                for (run, next_run) in zip(runs, runs[1:]):
                    split_run_into_paragraph(run, next_run, cell_paragraph)
                if runs:
                    split_run_into_paragraph(runs[-1], None, cell_paragraph)


def _embolden_word(word: str):
    bold_portion = math.floor(len(word) / 2)
    return word[:bold_portion], word[bold_portion:]
